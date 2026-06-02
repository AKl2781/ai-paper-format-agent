from __future__ import annotations

from docx import Document

from services.docx_analyzer import analyze_figure_tables


def make_document(paragraphs: list[str]) -> Document:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    return document


def check(paragraphs: list[str]) -> dict[str, object]:
    document = make_document(paragraphs)
    return analyze_figure_tables([p.text.strip() for p in document.paragraphs if p.text.strip()], document)


def assert_equal(name: str, actual: object, expected: object) -> None:
    if actual != expected:
        raise AssertionError(f"{name} FAIL expected={expected!r} actual={actual!r}")
    print(f"{name} PASS")


def main() -> None:
    decimal_figures = check(["图2.1 研究流程", "图2.2 实验结果", "图3.1 对照结果"])
    assert_equal("decimal_figure_numbers", decimal_figures["figure_numbers"], ["2.1", "2.2", "3.1"])
    assert_equal("decimal_figure_duplicates", decimal_figures["duplicate_figures"], [])

    hyphen_figures = check(["图2-1 研究流程", "图2-2 实验结果"])
    assert_equal("hyphen_figure_numbers", hyphen_figures["figure_numbers"], ["2.1", "2.2"])
    assert_equal("hyphen_figure_duplicates", hyphen_figures["duplicate_figures"], [])

    decimal_tables = check(["表3.1 样本分布", "表3.2 实验结果"])
    assert_equal("decimal_table_numbers", decimal_tables["table_numbers"], ["3.1", "3.2"])
    assert_equal("decimal_table_duplicates", decimal_tables["duplicate_tables"], [])

    letter_tables = check(["表A-1 变量定义", "表B-2 实验配置"])
    assert_equal("letter_table_numbers", letter_tables["table_numbers"], ["A.1", "B.2"])
    assert_equal("letter_table_duplicates", letter_tables["duplicate_tables"], [])

    print("COMPOSITE_NUMBERING PASS")


if __name__ == "__main__":
    main()
