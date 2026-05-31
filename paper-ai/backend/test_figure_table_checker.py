from __future__ import annotations

from docx import Document

from services.docx_analyzer import analyze_figure_tables


def make_document(paragraphs: list[str], *, table_count: int = 0) -> Document:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    for _ in range(table_count):
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "指标"
        table.cell(0, 1).text = "数值"
    return document


def check(paragraphs: list[str], *, table_count: int = 0) -> dict[str, object]:
    document = make_document(paragraphs, table_count=table_count)
    return analyze_figure_tables([p.text.strip() for p in document.paragraphs if p.text.strip()], document)


def assert_equal(name: str, actual: object, expected: object) -> None:
    if actual != expected:
        raise AssertionError(f"{name} FAIL expected={expected!r} actual={actual!r}")
    print(f"{name} PASS")


def main() -> None:
    figure_continuous = check(["图1 研究流程", "图2 实验结果"])
    assert_equal("figure_continuous", figure_continuous["figure_gaps"], [])

    table_continuous = check(["表1 样本分布", "表2 实验结果"])
    assert_equal("table_continuous", table_continuous["table_gaps"], [])

    figure_gap = check(["图1 研究流程", "图3 实验结果"])
    assert_equal("figure_gap", figure_gap["figure_gaps"], [2])

    table_gap = check(["表1 样本分布", "表3 实验结果"])
    assert_equal("table_gap", table_gap["table_gaps"], [2])

    duplicate_figure = check(["图1 研究流程", "图1 实验结果"])
    assert_equal("duplicate_figure", duplicate_figure["duplicate_figures"], [1])

    duplicate_table = check(["表1 样本分布", "表1 实验结果"])
    assert_equal("duplicate_table", duplicate_table["duplicate_tables"], [1])

    missing_figure_ref = check(["如图2所示，结果明显。", "图1 研究流程"])
    assert_equal("missing_referenced_figure", missing_figure_ref["missing_referenced_figures"], [2])

    missing_table_ref = check(["见表 3，结果明显。", "表1 样本分布"])
    assert_equal("missing_referenced_table", missing_table_ref["missing_referenced_tables"], [3])

    missing_table_caption = check(["正文只有一个 Word 表格。"], table_count=1)
    assert_equal("missing_table_caption", len(missing_table_caption["missing_table_captions"]), 1)

    print("FIGURE_TABLE_CHECKER PASS")


if __name__ == "__main__":
    main()
