from __future__ import annotations

import argparse
import csv
import os
import random
import shutil
import string
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent
CORPUS_DIR = BASE_DIR / "test_documents" / "generated"
MANIFEST_PATH = BASE_DIR / "test_documents" / "generated_manifest.csv"


SCENARIOS = [
    "clean",
    "messy",
    "references",
    "figures_tables",
    "reports",
    "template_mismatch",
    "medium",
]


PROFILE_COUNTS = {
    "quick": {
        "clean": 4,
        "messy": 4,
        "references": 4,
        "figures_tables": 4,
        "reports": 3,
        "template_mismatch": 3,
        "medium": 2,
    },
    "medium": {
        "clean": 10,
        "messy": 10,
        "references": 8,
        "figures_tables": 8,
        "reports": 6,
        "template_mismatch": 6,
        "medium": 5,
    },
    "stress": {
        "clean": 20,
        "messy": 20,
        "references": 16,
        "figures_tables": 16,
        "reports": 10,
        "template_mismatch": 10,
        "medium": 8,
    },
}


def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
    return len(data).to_bytes(4, "big") + chunk_type + data + crc.to_bytes(4, "big")


def write_png(path: Path, width: int, height: int, seed: int) -> None:
    random.seed(seed)
    rows = []
    for y in range(height):
        row = bytearray([0])
        for x in range(width):
            row.extend(
                [
                    (x * 7 + seed * 13) % 256,
                    (y * 11 + seed * 17) % 256,
                    ((x + y) * 5 + seed * 19) % 256,
                ]
            )
        rows.append(bytes(row))
    raw = b"".join(rows)
    header = width.to_bytes(4, "big") + height.to_bytes(4, "big") + bytes([8, 2, 0, 0, 0])
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + _png_chunk(b"IHDR", header)
        + _png_chunk(b"IDAT", zlib.compress(raw, level=6))
        + _png_chunk(b"IEND", b"")
    )


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=0)
    doc.add_paragraph("作者：测试用户")
    doc.add_paragraph("学院：数据与智能工程学院")
    doc.add_paragraph("专业：信息管理与信息系统")


def add_abstracts(doc: Document) -> None:
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本文围绕智能格式修复流程开展研究，构建面向论文文档的自动化处理链路，"
        "重点验证上传、分类、格式修复、评分、预览和下载等环节的稳定性。"
    )
    doc.add_paragraph("关键词：格式修复；文档处理；自动化测试；风险预检")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This paper validates an automated document formatting agent with local scoring, "
        "preview generation, and download workflows."
    )
    doc.add_paragraph("Keywords: formatting; document processing; regression testing")


def add_body(doc: Document, paragraphs: int, messy: bool = False) -> None:
    topic_words = ["模型", "数据", "流程", "规则", "质量", "模板", "审计", "性能"]
    for idx in range(1, paragraphs + 1):
        if idx % 18 == 1:
            doc.add_heading(f"{(idx // 18) + 1} 章节标题", level=1)
        if idx % 18 == 5:
            doc.add_heading(f"{(idx // 18) + 1}.{idx % 18} 小节标题", level=2)

        text = (
            f"第{idx}段围绕{random.choice(topic_words)}展开说明，"
            "用于覆盖正文段落、标题层级、中文标点、长句换行和局部重复表达。"
        )
        if messy and idx % 7 == 0:
            text = f"{(idx // 7) + 1}.混排标题：{text}"
        if messy and idx % 11 == 0:
            text += " C-51 附件占位内容需要被识别为模板残留。"
        doc.add_paragraph(text)


def add_references(doc: Document, count: int, variant: int) -> None:
    doc.add_heading("参考文献", level=1)
    for idx in range(1, count + 1):
        number = idx
        if variant == 1 and idx > 6:
            number = idx + 1
        if variant == 2 and idx == 8:
            number = 7
        if variant == 3 and idx == 10:
            doc.add_paragraph("未编号文献条目，作者. 文档测试方法研究[J]. 软件质量, 2026.")
            continue
        doc.add_paragraph(
            f"[{number}] 作者{idx}. 文档格式修复与自动化测试研究[J]. 软件工程实践, 2026, {idx}(1): 1-8."
        )


def add_tables(doc: Document, count: int, duplicate_number: bool = False) -> None:
    for idx in range(1, count + 1):
        table_no = 1 if duplicate_number and idx == 2 else idx
        doc.add_paragraph(f"表 {table_no} 测试指标统计")
        table = doc.add_table(rows=4, cols=4)
        table.style = "Table Grid"
        for r in range(4):
            for c in range(4):
                table.cell(r, c).text = f"R{r + 1}C{c + 1}-{idx}"


def add_figures(doc: Document, work_dir: Path, count: int, gap: bool = False) -> None:
    image_dir = work_dir / "_images"
    image_dir.mkdir(parents=True, exist_ok=True)
    for idx in range(1, count + 1):
        fig_no = idx + 1 if gap and idx >= 3 else idx
        image_path = image_dir / f"figure_{idx:03d}.png"
        write_png(image_path, 180, 90, seed=idx)
        doc.add_paragraph(f"图 {fig_no} 测试流程示意")
        doc.add_picture(str(image_path), width=Inches(2.8))


def build_document(scenario: str, index: int, output_path: Path) -> dict[str, str]:
    random.seed(f"{scenario}-{index}")
    doc = Document()

    title = f"{scenario.replace('_', ' ').title()} Test Thesis {index:03d}"
    add_title(doc, title)

    if scenario == "reports":
        doc.add_heading("实验目的", level=1)
        doc.add_paragraph("本报告用于验证非标准论文文档的分类和确认流程。")
        add_body(doc, 18, messy=True)
        expected_type = "lab_report"
        risks = "non_paper_confirmation"
    else:
        add_abstracts(doc)
        expected_type = "academic_paper"
        risks = "none"

    if scenario == "clean":
        add_body(doc, 28)
        add_references(doc, 8, variant=0)
    elif scenario == "messy":
        add_body(doc, 34, messy=True)
        add_references(doc, 8, variant=0)
        risks = "mixed_heading_body;template_residue"
    elif scenario == "references":
        add_body(doc, 26)
        add_references(doc, 12, variant=(index % 4))
        risks = "reference_numbering"
    elif scenario == "figures_tables":
        add_body(doc, 24)
        add_tables(doc, 4, duplicate_number=index % 2 == 0)
        add_figures(doc, output_path.parent, 4, gap=index % 2 == 1)
        add_references(doc, 6, variant=0)
        risks = "figure_table_numbering"
    elif scenario == "template_mismatch":
        add_body(doc, 20, messy=True)
        add_tables(doc, 2)
        add_references(doc, 5, variant=0)
        risks = "template_mismatch"
    elif scenario == "medium":
        add_body(doc, 160, messy=index % 2 == 0)
        add_tables(doc, 12, duplicate_number=index % 2 == 0)
        add_figures(doc, output_path.parent, 10, gap=index % 2 == 1)
        add_references(doc, 40, variant=index % 4)
        risks = "medium_scale"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {
        "case_id": output_path.stem,
        "file_name": str(output_path.relative_to(BASE_DIR / "test_documents")).replace(os.sep, "/"),
        "category": scenario,
        "document_type": expected_type,
        "template_file": "",
        "expected_classification_success": "true",
        "expected_format_success": "true",
        "expected_report": "true",
        "expected_preview": "true",
        "expected_download": "true",
        "expected_manual_review": "true" if risks != "none" else "false",
        "known_risks": risks,
        "notes": f"Generated {scenario} corpus sample.",
    }


def generate(profile: str, clean: bool) -> None:
    if clean and CORPUS_DIR.exists():
        shutil.rmtree(CORPUS_DIR)

    rows = []
    for scenario, count in PROFILE_COUNTS[profile].items():
        for index in range(1, count + 1):
            output_path = CORPUS_DIR / scenario / f"{scenario}_{index:03d}.docx"
            rows.append(build_document(scenario, index, output_path))

    MANIFEST_PATH.parent.mkdir(parents=True, exist_ok=True)
    with MANIFEST_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)

    image_dirs = list(CORPUS_DIR.rglob("_images"))
    for image_dir in image_dirs:
        shutil.rmtree(image_dir, ignore_errors=True)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate local DOCX regression corpus samples.")
    parser.add_argument("--profile", choices=PROFILE_COUNTS.keys(), default="quick")
    parser.add_argument("--clean", action="store_true", help="Remove generated corpus before writing.")
    args = parser.parse_args()
    generate(profile=args.profile, clean=args.clean)
    print(f"Generated profile={args.profile}")
    print(f"Corpus: {CORPUS_DIR}")
    print(f"Manifest: {MANIFEST_PATH}")


if __name__ == "__main__":
    main()
