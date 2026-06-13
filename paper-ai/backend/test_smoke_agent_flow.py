from __future__ import annotations

import io
import os
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from fastapi.testclient import TestClient

import services.language_reviewer as language_reviewer
from main import OUTPUT_DIR, app


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def make_template() -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    paragraph = doc.add_paragraph("模板正文样式")
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def assert_ok(name: str, condition: bool, detail: object = "") -> None:
    if not condition:
        raise AssertionError(f"{name} FAIL {detail}")
    suffix = f" {detail}" if detail else ""
    print(f"{name} PASS{suffix}")


def assert_pipeline_trace(result: dict[str, object]) -> None:
    trace = result.get("agent_trace")
    assert_ok("agent_trace_is_list", isinstance(trace, list) and bool(trace), trace)
    required_keys = {"step", "status", "duration_ms", "fallback_used", "message"}
    for index, item in enumerate(trace or []):
        assert_ok(f"agent_trace_item_{index}_shape", isinstance(item, dict) and required_keys.issubset(item.keys()), item)
        assert_ok(f"agent_trace_item_{index}_duration", isinstance(item.get("duration_ms"), int), item)
        assert_ok(f"agent_trace_item_{index}_fallback", isinstance(item.get("fallback_used"), bool), item)
    assert_ok("agent_trace_detail_compatible", isinstance(result.get("agent_trace_detail"), dict), result.get("agent_trace_detail"))


def main() -> None:
    client = TestClient(app)
    paper = make_docx(
        [
            "高中生压力管理研究",
            "摘要：本文围绕高中生压力管理展开分析。",
            "关键词：压力管理；高中生；教育研究",
            "1.1研究背景",
            "当前高中生面对学业与生活压力。",
            "1.2 [高中生压力大的原因] 高中生压力来源包括学业任务、人际关系以及升学期待。",
            "2.研究方法",
            "本文采用文献分析方法。",
            "4.结语：本文认为需要从学校、家庭和学生自身三个方面改进压力管理。C-51",
            "参考文献",
            "[1] 张三. 高中生压力研究[J]. 教育研究, 2024.",
        ]
    )

    response = client.post("/document/classify", files={"paper": ("smoke-paper.docx", paper, DOCX_MIME)})
    data = response.json()
    assert_ok("classify_endpoint", response.status_code == 200, response.status_code)
    assert_ok("standard_paper_not_unknown", data.get("document_type") == "academic_paper", data.get("document_type"))

    response = client.post(
        "/agent/run",
        data={"mode": "local"},
        files={"paper": ("smoke-local.docx", paper, DOCX_MIME)},
    )
    local = response.json()
    assert_ok("local_without_template_flow", response.status_code == 200 and local.get("status") == "ok", local.get("status"))
    breakdown = local.get("score_breakdown", {})
    assert_ok("local_ai_score_null", breakdown.get("ai_score") is None, breakdown)
    assert_ok("local_ai_used_false", breakdown.get("ai_used") is False, breakdown)
    local_output = OUTPUT_DIR / local["filename"]
    assert_ok("local_output_file_created", local_output.exists() and local_output.stat().st_size > 0, local_output.name)
    assert_ok("local_report_created", bool(local.get("modification_report")), local["modification_report"]["change_counts"])
    assert_ok("top_level_reference_check", isinstance(local.get("reference_check"), dict), local.get("reference_check"))
    assert_ok("top_level_figure_table_check", isinstance(local.get("figure_table_check"), dict), local.get("figure_table_check"))
    assert_pipeline_trace(local)
    local_preview = client.get(f"/preview/{local['filename']}")
    assert_ok("local_preview_endpoint", local_preview.status_code == 200 and bool(local_preview.json().get("html")), local_preview.status_code)
    local_download = client.get(local["download_url"])
    assert_ok("local_download_endpoint", local_download.status_code == 200 and len(local_download.content) > 0, local_download.status_code)

    response = client.post(
        "/agent/run",
        data={"mode": "local"},
        files={
            "paper": ("smoke-template.docx", paper, DOCX_MIME),
            "template": ("smoke-template-file.docx", make_template(), DOCX_MIME),
        },
    )
    templated = response.json()
    assert_ok("local_with_template_flow", response.status_code == 200 and templated.get("status") == "ok", templated.get("status"))
    templated_output = OUTPUT_DIR / templated["filename"]
    assert_ok("template_output_file_created", templated_output.exists() and templated_output.stat().st_size > 0, templated_output.name)
    templated_text = "\n".join(p.text for p in Document(templated_output).paragraphs)
    assert_ok("no_unsupported_operand_error", "unsupported operand type" not in str(templated), "")
    assert_ok("c51_removed", "C-51" not in templated_text, "")
    assert_ok("mixed_heading_split", "4.结语\n本文认为" in templated_text, "")
    templated_download = client.get(templated["download_url"])
    assert_ok("template_download_endpoint", templated_download.status_code == 200 and len(templated_download.content) > 0, templated_download.status_code)

    def fail_ai(_path: Path) -> dict[str, object]:
        raise RuntimeError("simulated ai failure")

    os.environ["DEEPSEEK_API_KEY"] = "fake-key"
    os.environ.pop("OPENAI_API_KEY", None)
    language_reviewer.call_ai_language_review = fail_ai
    response = client.post(
        "/agent/run",
        data={"mode": "ai"},
        files={"paper": ("smoke-ai.docx", paper, DOCX_MIME)},
    )
    ai = response.json()
    assert_ok("ai_failure_fallback_flow", response.status_code == 200 and ai.get("status") == "ok", ai.get("status"))
    assert_ok("ai_fallback_output_file_created", (OUTPUT_DIR / ai["filename"]).exists(), ai.get("filename"))
    assert_ok("ai_fallback_ai_used_false", ai.get("score_breakdown", {}).get("ai_used") is False, ai.get("score_breakdown"))
    ai_download = client.get(ai["download_url"])
    assert_ok("ai_fallback_download_endpoint", ai_download.status_code == 200 and len(ai_download.content) > 0, ai_download.status_code)
    print("SMOKE PASS")


if __name__ == "__main__":
    main()
