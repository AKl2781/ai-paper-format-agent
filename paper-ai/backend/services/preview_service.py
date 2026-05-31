from __future__ import annotations

import re
from html import escape
from pathlib import Path
from typing import Any

from docx import Document


def build_docx_preview(path: Path) -> dict[str, str]:
    document = Document(path)
    title = find_title(document)
    html_parts: list[str] = []
    in_references = False
    reference_parts: list[str] = []
    has_content = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if is_reference_heading(text):
            if reference_parts:
                html_parts.append(render_reference_section(reference_parts))
                reference_parts = []
            in_references = True
            html_parts.append(render_paragraph(paragraph, text, is_first=False, extra_class="docx-reference-heading"))
            has_content = True
            continue

        if in_references:
            reference_parts.append(render_paragraph(paragraph, text, is_first=False, extra_class="docx-reference-item"))
            has_content = True
            continue

        html_parts.append(render_paragraph(paragraph, text, is_first=not has_content))
        has_content = True

    if reference_parts:
        html_parts.append(render_reference_section(reference_parts))

    for table in document.tables:
        html_parts.append(render_table(table))

    return {
        "title": title,
        "html": "\n".join(html_parts) or '<p class="docx-paragraph">未能读取到可预览内容。</p>',
    }


def render_paragraph(paragraph, text: str, *, is_first: bool, extra_class: str = "") -> str:
    tag = paragraph_tag(paragraph, text, is_first=is_first)
    content = render_runs(paragraph)
    align = alignment_class(paragraph.alignment)
    level_class = heading_level_class(tag)
    classes = " ".join(item for item in ["docx-paragraph", align, level_class, extra_class] if item)
    return f'<{tag} class="{classes}">{content}</{tag}>'


def render_reference_section(reference_parts: list[str]) -> str:
    return f'<section class="docx-references">{"".join(reference_parts)}</section>'


def find_title(document) -> str:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text[:120]
    return "论文预览"


def paragraph_tag(paragraph, text: str, *, is_first: bool = False) -> str:
    style_name = paragraph.style.name if paragraph.style else ""
    if is_first or style_name.startswith("Title") or style_name.startswith("论文标题"):
        return "h1"
    if style_name.startswith("Heading 2") or style_name.startswith("标题 2") or is_second_level_heading(text):
        return "h3"
    if style_name.startswith("Heading 3") or style_name.startswith("标题 3") or is_third_level_heading(text):
        return "h4"
    if style_name.startswith("Heading 1") or style_name.startswith("标题 1") or is_main_heading(text):
        return "h2"
    return "p"


def heading_level_class(tag: str) -> str:
    if tag in {"h1", "h2", "h3", "h4"}:
        return f"docx-{tag}"
    return "docx-body"


def is_main_heading(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    return (
        compact in {"摘要", "关键词", "关键字", "引言", "绪论", "结论", "参考文献"}
        or compact.lower() in {"abstract", "keywords", "introduction", "conclusion", "references", "bibliography"}
        or bool(re.match(r"^第[一二三四五六七八九十]+章", compact))
        or bool(re.match(r"^[一二三四五六七八九十]+[、.．]\S{1,30}$", compact))
        or bool(re.match(r"^\d+[、.．](?!\d)\S{1,30}$", compact))
    )


def is_second_level_heading(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    return bool(re.match(r"^\d+\.\d+\s*\S{1,30}$", compact))


def is_third_level_heading(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    return bool(re.match(r"^\d+\.\d+\.\d+\s*\S{1,30}$", compact))


def is_reference_heading(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    return compact in {"参考文献"} or compact.lower() in {"references", "bibliography"}


def render_runs(paragraph) -> str:
    if not paragraph.runs:
        return escape(paragraph.text.strip())
    parts: list[str] = []
    for run in paragraph.runs:
        text = escape(run.text)
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    return "".join(parts).strip() or escape(paragraph.text.strip())


def render_table(table) -> str:
    rows = []
    for row_index, row in enumerate(table.rows):
        cell_tag = "th" if row_index == 0 else "td"
        cells = "".join(f"<{cell_tag}>{escape(cell.text.strip())}</{cell_tag}>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return f'<div class="docx-table-wrap"><table class="docx-table">{"".join(rows)}</table></div>'


def alignment_class(alignment: Any) -> str:
    name = str(alignment)
    if "CENTER" in name:
        return "align-center"
    if "RIGHT" in name:
        return "align-right"
    if "JUSTIFY" in name:
        return "align-justify"
    return "align-left"
