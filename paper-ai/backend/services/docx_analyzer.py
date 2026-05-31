from __future__ import annotations

import re
from pathlib import Path
from statistics import mean
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


LOCAL_WEIGHTS = {
    "structure": 0.18,
    "heading_format": 0.14,
    "body_font": 0.12,
    "spacing_indent": 0.16,
    "page_margin": 0.14,
    "references": 0.12,
    "repeat_risk": 0.14,
}

AI_WEIGHTS = {
    "language": 0.25,
    "fluency": 0.25,
    "logic": 0.25,
    "academic_expression": 0.25,
}


def analyze_docx(
    path: Path,
    *,
    template_path: Path | None = None,
    repeat_score: int | None = None,
    ai_scores: dict[str, int] | None = None,
) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    headings = [item for item in paragraphs if is_section_heading(item)]
    table_row_count = sum(len(table.rows) for table in document.tables)
    reference_check = analyze_references(paragraphs)
    figure_table_check = analyze_figure_tables(paragraphs, document)
    local_breakdown = build_local_breakdown(document, paragraphs, headings, template_path, repeat_score, reference_check)
    ai_breakdown = build_ai_breakdown(ai_scores) if ai_scores else []
    score_breakdown = build_score_breakdown(local_breakdown, ai_breakdown)

    return {
        "paragraph_count": len(paragraphs),
        "character_count": len(text),
        "heading_count": max(count_styled_headings(document), len(headings)),
        "table_row_count": table_row_count,
        "preview": text[:1200],
        "reference_check": reference_check,
        "figure_table_check": figure_table_check,
        "report": {
            "score": score_breakdown["final_score"],
            "summary": make_summary(score_breakdown),
            "breakdown": [*local_breakdown, *ai_breakdown],
            "local_breakdown": local_breakdown,
            "ai_breakdown": ai_breakdown,
            "score_breakdown": score_breakdown,
            "issues": flatten_issues([*local_breakdown, *ai_breakdown]),
            "recommendations": build_recommendations(table_row_count, template_path),
        },
    }


def extract_paragraph_text(path: Path) -> list[str]:
    document = Document(path)
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


def build_local_breakdown(
    document,
    paragraphs: list[str],
    headings: list[str],
    template_path: Path | None,
    repeat_score: int | None,
    reference_check: dict[str, Any],
) -> list[dict[str, Any]]:
    return [
        score_structure(paragraphs, headings),
        score_heading_format(document, headings),
        score_body_font(document),
        score_spacing_indent(document),
        score_page_margin(document),
        score_references(reference_check),
        dimension("repeat_risk", "重复风险预检", repeat_score if repeat_score is not None else 88, ["相似度预检不能替代正式查重。"], "local"),
    ]


def build_ai_breakdown(ai_scores: dict[str, int] | None) -> list[dict[str, Any]]:
    if not ai_scores:
        return []
    return [
        dimension("language", "语言规范性", ai_scores.get("language", 82), ["AI 已审校语言规范性，仍建议人工复读关键段落。"], "ai"),
        dimension("fluency", "表达流畅度", ai_scores.get("fluency", 82), ["AI 已评估表达流畅度。"], "ai"),
        dimension("logic", "论文逻辑性", ai_scores.get("logic", 80), ["AI 已评估论证衔接，但不会补写事实。"], "ai"),
        dimension("academic_expression", "学术表达质量", ai_scores.get("academic_expression", 82), ["AI 已评估学术表达质量。"], "ai"),
    ]


def build_score_breakdown(local_breakdown: list[dict[str, Any]], ai_breakdown: list[dict[str, Any]]) -> dict[str, Any]:
    local_score = weighted(local_breakdown, LOCAL_WEIGHTS)
    if not ai_breakdown:
        return {
            "local_score": local_score,
            "ai_score": None,
            "final_score": local_score,
            "ai_used": False,
            "ai_added_value": [],
        }
    ai_score = weighted(ai_breakdown, AI_WEIGHTS)
    final_score = max(local_score, min(96, round(local_score * 0.78 + ai_score * 0.22)))
    return {
        "local_score": local_score,
        "ai_score": ai_score,
        "final_score": final_score,
        "ai_used": True,
        "ai_added_value": build_ai_added_value(ai_breakdown, local_score, final_score),
    }


def weighted(items: list[dict[str, Any]], weights: dict[str, float]) -> int:
    return max(0, min(96, round(sum(item["score"] * weights[item["key"]] for item in items))))


def build_ai_added_value(ai_breakdown: list[dict[str, Any]], local_score: int, final_score: int) -> list[str]:
    values = [f"{item['label']}：{item['score']} 分" for item in ai_breakdown]
    values.append(f"AI增强使最终评分相对本地模式提升 {max(0, final_score - local_score)} 分。")
    return values


def score_structure(paragraphs: list[str], headings: list[str]) -> dict[str, Any]:
    text = "\n".join(paragraphs).lower()
    checks = [
        ("论文标题", bool(paragraphs and 4 <= len(paragraphs[0]) <= 80)),
        ("摘要", contains_any(text, ["摘要", "abstract"])),
        ("关键词", contains_any(text, ["关键词", "关键字", "keywords"])),
        ("章节标题", bool(headings)),
        ("参考文献", contains_any(text, ["参考文献", "references", "bibliography"])),
        ("正文段落", len(paragraphs) >= 5),
    ]
    passed = sum(1 for _, ok in checks if ok)
    return dimension("structure", "结构完整度", round(55 + passed / len(checks) * 40), [f"缺少或未清晰识别：{name}" for name, ok in checks if not ok], "local")


def score_heading_format(document, headings: list[str]) -> dict[str, Any]:
    if not headings:
        return dimension("heading_format", "标题格式", 66, ["未识别到清晰章节标题。"], "local")
    styled = sum(1 for p in document.paragraphs if is_section_heading(p.text.strip()) and is_heading_style(p.style.name if p.style else ""))
    ratio = styled / max(len(headings), 1)
    return dimension("heading_format", "标题格式", round(68 + ratio * 24), [] if ratio >= 0.8 else ["部分章节标题未套用 Word 标题样式。"], "local")


def score_body_font(document) -> dict[str, Any]:
    body = get_body_paragraphs(document)
    if not body:
        return dimension("body_font", "正文字体", 72, ["未识别到足够正文段落。"], "local")
    ok = 0
    for paragraph in body:
        sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
        if not sizes or 10.5 <= mean(sizes) <= 12.5:
            ok += 1
    ratio = ok / len(body)
    return dimension("body_font", "正文字体", round(72 + ratio * 18), [] if ratio >= 0.8 else ["较多正文段落字号不在常见论文正文范围。"], "local")


def score_spacing_indent(document) -> dict[str, Any]:
    body = get_body_paragraphs(document)
    if not body:
        return dimension("spacing_indent", "行距缩进", 70, ["未识别到足够正文段落。"], "local")
    line_ok = sum(1 for p in body if line_spacing_ok(p)) / len(body)
    indent_ok = sum(1 for p in body if first_line_indent_ok(p)) / len(body)
    align_ok = sum(1 for p in body if paragraph_alignment_ok(p)) / len(body)
    issues = []
    if line_ok < 0.8:
        issues.append("正文行距仍需复核。")
    if indent_ok < 0.8:
        issues.append("正文首行缩进仍需复核。")
    if align_ok < 0.8:
        issues.append("正文对齐方式不够统一。")
    return dimension("spacing_indent", "行距缩进", round(58 + line_ok * 16 + indent_ok * 16 + align_ok * 8), issues, "local")


def score_page_margin(document) -> dict[str, Any]:
    section = document.sections[0]
    checks = [
        near(section.top_margin.cm, 2.54, 0.25),
        near(section.bottom_margin.cm, 2.54, 0.25),
        near(section.left_margin.cm, 3.18, 0.35),
        near(section.right_margin.cm, 3.18, 0.35),
    ]
    ratio = sum(1 for ok in checks if ok) / len(checks)
    return dimension("page_margin", "页边距", round(65 + ratio * 27), [] if ratio >= 0.75 else ["页边距与常见论文规范差异较大。"], "local")


def score_references(reference_check: dict[str, Any]) -> dict[str, Any]:
    if not reference_check["has_reference_section"]:
        return dimension("references", "参考文献基础格式", 58, ["未识别到参考文献部分。"], "local")
    if reference_check["reference_count"] == 0:
        return dimension("references", "参考文献基础格式", 62, ["参考文献标题存在，但未识别到条目。"], "local")

    issues = reference_check["issues"]
    formatted_ratio = min(1, len(reference_check["reference_numbers"]) / max(reference_check["reference_count"], 1))
    score = round(70 + formatted_ratio * 18)
    score -= len(reference_check["numbering_gaps"]) * 4
    score -= len(reference_check["duplicate_reference_numbers"]) * 6
    score -= len(reference_check["missing_reference_numbers"]) * 5
    score -= min(12, len(reference_check["uncited_reference_numbers"]) * 2)
    return dimension("references", "参考文献基础格式", score, issues, "local")


def analyze_references(paragraphs: list[str]) -> dict[str, Any]:
    ref_index = next((i for i, text in enumerate(paragraphs) if is_reference_heading(text)), -1)
    if ref_index < 0:
        return build_reference_check(
            has_reference_section=False,
            reference_count=0,
            citation_numbers=[],
            reference_numbers=[],
            duplicate_reference_numbers=[],
            numbering_gaps=[],
            missing_reference_numbers=[],
            uncited_reference_numbers=[],
            issues=["未识别到参考文献部分。"],
        )

    body_paragraphs = paragraphs[:ref_index]
    reference_items = paragraphs[ref_index + 1 :]
    citation_numbers = extract_citation_numbers(body_paragraphs)
    parsed_reference_numbers = [parse_reference_number(item) for item in reference_items]
    reference_numbers = [number for number in parsed_reference_numbers if number is not None]
    duplicate_reference_numbers = repeated_numbers(reference_numbers)
    unique_reference_numbers = sorted(set(reference_numbers))
    unique_citation_numbers = sorted(set(citation_numbers))
    numbering_gaps = find_numbering_gaps(unique_reference_numbers)
    missing_reference_numbers = sorted(set(unique_citation_numbers) - set(unique_reference_numbers))
    uncited_reference_numbers = sorted(set(unique_reference_numbers) - set(unique_citation_numbers))
    formatted_count = sum(1 for number in parsed_reference_numbers if number is not None)
    formatted_reference_ratio = formatted_count / max(len(reference_items), 1)
    issues = build_reference_issues(
        has_items=bool(reference_items),
        formatted_reference_ratio=formatted_reference_ratio,
        duplicate_reference_numbers=duplicate_reference_numbers,
        numbering_gaps=numbering_gaps,
        missing_reference_numbers=missing_reference_numbers,
        uncited_reference_numbers=uncited_reference_numbers,
    )

    return build_reference_check(
        has_reference_section=True,
        reference_count=len(reference_items),
        citation_numbers=unique_citation_numbers,
        reference_numbers=unique_reference_numbers,
        duplicate_reference_numbers=duplicate_reference_numbers,
        numbering_gaps=numbering_gaps,
        missing_reference_numbers=missing_reference_numbers,
        uncited_reference_numbers=uncited_reference_numbers,
        issues=issues,
    )


def build_reference_check(
    *,
    has_reference_section: bool,
    reference_count: int,
    citation_numbers: list[int],
    reference_numbers: list[int],
    duplicate_reference_numbers: list[int],
    numbering_gaps: list[int],
    missing_reference_numbers: list[int],
    uncited_reference_numbers: list[int],
    issues: list[str],
) -> dict[str, Any]:
    return {
        "has_reference_section": has_reference_section,
        "reference_count": reference_count,
        "citation_count": len(citation_numbers),
        "reference_numbers": reference_numbers,
        "citation_numbers": citation_numbers,
        "missing_reference_numbers": missing_reference_numbers,
        "uncited_reference_numbers": uncited_reference_numbers,
        "duplicate_reference_numbers": duplicate_reference_numbers,
        "numbering_gaps": numbering_gaps,
        "issues": issues,
    }


def extract_citation_numbers(paragraphs: list[str]) -> list[int]:
    numbers: list[int] = []
    for text in paragraphs:
        for group in re.findall(r"\[(\d+(?:\s*[-,，、]\s*\d+)*)\]", text):
            numbers.extend(expand_citation_group(group))
    return sorted(numbers)


def expand_citation_group(group: str) -> list[int]:
    numbers: list[int] = []
    for part in re.split(r"[,，、]\s*", group):
        compact = part.strip()
        if not compact:
            continue
        range_match = re.match(r"^(\d+)\s*-\s*(\d+)$", compact)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            if start <= end and end - start <= 20:
                numbers.extend(range(start, end + 1))
            continue
        if compact.isdigit():
            numbers.append(int(compact))
    return numbers


def parse_reference_number(text: str) -> int | None:
    match = re.match(r"^\s*(?:\[(\d+)\]|(\d+)[.．、])", text)
    if not match:
        return None
    return int(match.group(1) or match.group(2))


def repeated_numbers(numbers: list[int]) -> list[int]:
    return sorted({number for number in numbers if numbers.count(number) > 1})


def find_numbering_gaps(numbers: list[int]) -> list[int]:
    if len(numbers) < 2:
        return []
    expected = set(range(min(numbers), max(numbers) + 1))
    return sorted(expected - set(numbers))


def build_reference_issues(
    *,
    has_items: bool,
    formatted_reference_ratio: float,
    duplicate_reference_numbers: list[int],
    numbering_gaps: list[int],
    missing_reference_numbers: list[int],
    uncited_reference_numbers: list[int],
) -> list[str]:
    issues: list[str] = []
    if not has_items:
        issues.append("参考文献标题存在，但未识别到条目。")
    if has_items and formatted_reference_ratio < 0.75:
        issues.append("参考文献编号或条目格式不够统一。")
    if duplicate_reference_numbers:
        issues.append(f"参考文献存在重复编号：{format_number_list(duplicate_reference_numbers)}。")
    if numbering_gaps:
        issues.append(f"参考文献编号不连续，缺少：{format_number_list(numbering_gaps)}。")
    if missing_reference_numbers:
        issues.append(f"正文引用编号未在文末参考文献中找到：{format_number_list(missing_reference_numbers)}。")
    if uncited_reference_numbers:
        issues.append(f"文末参考文献未在正文中引用：{format_number_list(uncited_reference_numbers)}。")
    return issues


def format_number_list(numbers: list[int]) -> str:
    return "、".join(str(number) for number in numbers)


def analyze_figure_tables(paragraphs: list[str], document) -> dict[str, Any]:
    figure_numbers = extract_caption_numbers(paragraphs, "figure")
    table_numbers = extract_caption_numbers(paragraphs, "table")
    referenced_figures = extract_referenced_numbers(paragraphs, "figure")
    referenced_tables = extract_referenced_numbers(paragraphs, "table")
    unique_figures = sorted(set(figure_numbers))
    unique_tables = sorted(set(table_numbers))
    unique_referenced_figures = sorted(set(referenced_figures))
    unique_referenced_tables = sorted(set(referenced_tables))
    duplicate_figures = repeated_numbers(figure_numbers)
    duplicate_tables = repeated_numbers(table_numbers)
    figure_gaps = find_numbering_gaps(unique_figures)
    table_gaps = find_numbering_gaps(unique_tables)
    missing_referenced_figures = sorted(set(unique_referenced_figures) - set(unique_figures))
    missing_referenced_tables = sorted(set(unique_referenced_tables) - set(unique_tables))
    missing_figure_captions = detect_missing_figure_captions(document, len(unique_figures))
    missing_table_captions = detect_missing_table_captions(document, len(unique_tables))
    issues = build_figure_table_issues(
        figure_gaps=figure_gaps,
        table_gaps=table_gaps,
        duplicate_figures=duplicate_figures,
        duplicate_tables=duplicate_tables,
        missing_figure_captions=missing_figure_captions,
        missing_table_captions=missing_table_captions,
        missing_referenced_figures=missing_referenced_figures,
        missing_referenced_tables=missing_referenced_tables,
    )
    return {
        "figure_numbers": unique_figures,
        "table_numbers": unique_tables,
        "figure_gaps": figure_gaps,
        "table_gaps": table_gaps,
        "duplicate_figures": duplicate_figures,
        "duplicate_tables": duplicate_tables,
        "missing_figure_captions": missing_figure_captions,
        "missing_table_captions": missing_table_captions,
        "missing_referenced_figures": missing_referenced_figures,
        "missing_referenced_tables": missing_referenced_tables,
        "issues": issues,
    }


def extract_caption_numbers(paragraphs: list[str], kind: str) -> list[int]:
    numbers: list[int] = []
    for text in paragraphs:
        number = parse_caption_number(text, kind)
        if number is not None:
            numbers.append(number)
    return numbers


def parse_caption_number(text: str, kind: str) -> int | None:
    compact = text.strip()
    if kind == "figure":
        match = re.match(r"^(?:图\s*|Figure\s+)(\d+)(?:[-.．]\d+)?(?:\s|[：:、.．-]|$)", compact, re.IGNORECASE)
    else:
        match = re.match(r"^(?:表\s*|Table\s+)(\d+)(?:[-.．]\d+)?(?:\s|[：:、.．-]|$)", compact, re.IGNORECASE)
    return int(match.group(1)) if match else None


def extract_referenced_numbers(paragraphs: list[str], kind: str) -> list[int]:
    numbers: list[int] = []
    for text in paragraphs:
        if parse_caption_number(text, kind) is not None:
            continue
        patterns = figure_reference_patterns() if kind == "figure" else table_reference_patterns()
        for pattern in patterns:
            numbers.extend(int(match) for match in re.findall(pattern, text, re.IGNORECASE))
    return numbers


def figure_reference_patterns() -> list[str]:
    return [
        r"(?:如|见|参见|详见|由|从|根据)\s*图\s*(\d+)",
        r"图\s*(\d+)\s*所示",
        r"\bFigure\s+(\d+)\b",
    ]


def table_reference_patterns() -> list[str]:
    return [
        r"(?:如|见|参见|详见|由|从|根据)\s*表\s*(\d+)",
        r"表\s*(\d+)\s*所示",
        r"\bTable\s+(\d+)\b",
    ]


def detect_missing_figure_captions(document, caption_count: int) -> list[str]:
    image_count = count_inline_images(document)
    if image_count > caption_count:
        return [f"检测到约 {image_count} 个图片对象，但仅识别到 {caption_count} 个图题。"]
    return []


def detect_missing_table_captions(document, caption_count: int) -> list[str]:
    table_count = len(document.tables)
    if table_count > caption_count:
        return [f"检测到 {table_count} 个 Word 表格，但仅识别到 {caption_count} 个表题。"]
    return []


def count_inline_images(document) -> int:
    return sum(1 for _ in document.element.body.iter() if str(_.tag).endswith("}drawing") or str(_.tag).endswith("}pict"))


def build_figure_table_issues(
    *,
    figure_gaps: list[int],
    table_gaps: list[int],
    duplicate_figures: list[int],
    duplicate_tables: list[int],
    missing_figure_captions: list[str],
    missing_table_captions: list[str],
    missing_referenced_figures: list[int],
    missing_referenced_tables: list[int],
) -> list[str]:
    issues: list[str] = []
    if figure_gaps:
        issues.append(f"图编号不连续，缺少：{format_number_list(figure_gaps)}。")
    if table_gaps:
        issues.append(f"表编号不连续，缺少：{format_number_list(table_gaps)}。")
    if duplicate_figures:
        issues.append(f"图编号存在重复：{format_number_list(duplicate_figures)}。")
    if duplicate_tables:
        issues.append(f"表编号存在重复：{format_number_list(duplicate_tables)}。")
    issues.extend(missing_figure_captions)
    issues.extend(missing_table_captions)
    if missing_referenced_figures:
        issues.append(f"正文引用的图号未找到对应图题：{format_number_list(missing_referenced_figures)}。")
    if missing_referenced_tables:
        issues.append(f"正文引用的表号未找到对应表题：{format_number_list(missing_referenced_tables)}。")
    return issues


def dimension(key: str, label: str, score: int, issues: list[str], group: str) -> dict[str, Any]:
    bounded = max(0, min(96, int(score)))
    return {"key": key, "label": label, "score": bounded, "group": group, "status": score_status(bounded), "issues": issues}


def count_styled_headings(document) -> int:
    return sum(1 for p in document.paragraphs if p.text.strip() and is_heading_style(p.style.name if p.style else ""))


def get_body_paragraphs(document) -> list[Any]:
    body = []
    title_seen = False
    in_references = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if not title_seen:
            title_seen = True
            continue
        if is_reference_heading(text):
            in_references = True
            continue
        if in_references or is_section_heading(text) or text.startswith(("摘要", "关键词", "关键字", "Abstract", "Keywords")):
            continue
        body.append(paragraph)
    return body


def is_heading_style(style_name: str) -> bool:
    return style_name.startswith("Heading") or style_name.startswith("标题")


def is_section_heading(text: str) -> bool:
    compact = text.strip()
    if not compact or len(compact) > 45:
        return False
    compact = compact.rstrip("：:")
    if compact in {"摘要", "关键词", "关键字", "引言", "绪论", "结论", "结语", "致谢", "参考文献"}:
        return True
    if compact.lower() in {"abstract", "keywords", "introduction", "conclusion", "references", "bibliography"}:
        return True
    return is_probable_heading_prefix(compact)


def split_mixed_heading_text(text: str) -> tuple[str, str] | None:
    compact = text.strip()
    if len(compact) < 12:
        return None

    colon_match = re.match(r"^(.{2,45}?)[：:]\s*(\S.{5,})$", compact)
    if colon_match and is_probable_heading_prefix(colon_match.group(1).strip()):
        return colon_match.group(1).strip(), colon_match.group(2).strip()

    bracket_match = re.match(
        r"^((?:\d+\.\d+(?:\.\d+){0,2}|[一二三四五六七八九十]+[、.．]|第[一二三四五六七八九十\d]+[章节])\s*[\[【（(][^\]】）)]{2,35}[\]】）)])\s*(\S.{5,})$",
        compact,
    )
    if bracket_match and is_probable_heading_prefix(bracket_match.group(1).strip()):
        return bracket_match.group(1).strip(), bracket_match.group(2).strip()

    return None


def is_probable_heading_prefix(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    if len(compact) > 45:
        return False
    if re.match(r"^(第[一二三四五六七八九十\d]+[章节]|[一二三四五六七八九十]+[、.．])\s*\S+", compact):
        return True
    if re.match(r"^\d+[、.．]\s*\S+", compact):
        return True
    if re.match(r"^\d+\.\d+(?:\.\d+){0,2}\s*\S+", compact):
        return True
    return bool(re.match(r"^\d{1,2}\s*(引言|绪论|结论|结语|总结|参考文献|研究背景|研究方法|结果与分析|讨论)$", compact))


def guess_heading_level(text: str) -> int:
    if re.match(r"^[0-9]+\.[0-9]+\.[0-9]+", text):
        return 3
    if re.match(r"^[0-9]+\.[0-9]+", text):
        return 2
    return 1


def is_reference_heading(text: str) -> bool:
    compact = text.strip().rstrip("：:")
    return compact in {"参考文献"} or compact.lower() in {"references", "bibliography"}


def line_spacing_ok(paragraph) -> bool:
    value = paragraph.paragraph_format.line_spacing
    return isinstance(value, float) and 1.35 <= value <= 1.75


def first_line_indent_ok(paragraph) -> bool:
    value = paragraph.paragraph_format.first_line_indent
    return value is not None and 0.55 <= value.cm <= 0.95


def paragraph_alignment_ok(paragraph) -> bool:
    return paragraph.alignment in {WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT}


def near(value: float, target: float, tolerance: float) -> bool:
    return abs(value - target) <= tolerance


def contains_any(text: str, terms: list[str]) -> bool:
    return any(term.lower() in text for term in terms)


def flatten_issues(items: list[dict[str, Any]]) -> list[dict[str, str]]:
    issues = []
    for item in items:
        for issue in item["issues"]:
            issues.append({"severity": severity_from_score(item["score"]), "title": item["label"], "detail": issue})
    return issues


def score_status(score: int) -> str:
    if score >= 90:
        return "良好"
    if score >= 78:
        return "可提交前复核"
    if score >= 65:
        return "需优化"
    return "风险较高"


def severity_from_score(score: int) -> str:
    if score < 65:
        return "high"
    if score < 80:
        return "medium"
    return "low"


def make_summary(score_breakdown: dict[str, Any]) -> str:
    if score_breakdown["ai_used"]:
        return "已完成本地格式修复与 AI 增强审校，仍建议人工复核关键内容。"
    return "已完成本地格式修复与基础预检，未启用 AI 增强评分。"


def build_recommendations(table_row_count: int, template_path: Path | None) -> list[str]:
    items = [
        "人工核对封面、页眉页脚、脚注、图片题注、公式编号等复杂元素。",
        "检查参考文献真实性、引用对应关系和学校指定格式。",
        "使用 Word 打开最终文件，确认分页和目录层级是否符合提交要求。",
    ]
    if table_row_count:
        items.append("检测到表格，请确认表题、编号、单位和表头是否完整。")
    if not template_path:
        items.append("未上传模板，本次按通用论文规范估算。")
    return items
