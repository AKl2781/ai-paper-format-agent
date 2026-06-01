from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .docx_analyzer import guess_heading_level, is_section_heading, split_mixed_heading_text
from .template_extractor import extract_template_profile, safe_margin, safe_size, safe_spacing


def apply_paper_format(source: Path, output: Path, template: Path | None = None) -> list[str]:
    document = Document(source)
    template_profile = extract_template_profile(template) if template else None
    applied: list[str] = []

    applied.extend(clean_document_text(document))
    applied.extend(split_mixed_heading_paragraphs(document))
    apply_page_setup(document, template_profile)
    apply_normal_style(document, template_profile)
    applied.extend(split_long_paragraphs(document))
    apply_paragraph_styles(document, template_profile)
    apply_table_styles(document)

    applied.append(f"已读取模板并套用关键格式：{template.name}" if template_profile and template else "未上传模板，已使用通用论文格式规则。")
    applied.append("已统一正文行距、首行缩进、段前段后和标题层级。")
    if document.tables:
        applied.append("已统一表格内文字对齐和行距。")

    document.save(output)
    return applied


def clean_document_text(document) -> list[str]:
    changes: list[str] = []
    removed_placeholders = 0
    removed_style_notes = 0
    removed_template_codes = 0
    normalized_labels = 0
    keyword_suggestion = suggest_keywords("\n".join(p.text for p in document.paragraphs))

    for paragraph in list(document.paragraphs):
        original = paragraph.text
        cleaned = original.strip()
        cleaned, count = re.subn(r"\[[^\]]*(宋体|黑体|楷体|字号|行距|居中|左对齐|右对齐|C-\d+|B-\d+)[^\]]*\]", "", cleaned)
        removed_style_notes += count
        cleaned, count = re.subn(r"[（(【\[]?\b[BC]-\d{1,3}\b[）)】\]]?", "", cleaned)
        removed_template_codes += count
        cleaned, count = re.subn(r"[“”\"']?(作者姓名|学生姓名|论文题目|学校名称|学院名称|专业班级|学号|指导教师)[“”\"']?", "", cleaned)
        removed_placeholders += count
        cleaned, count = re.subn(r"^\s*(中文摘要|摘要)\s*[:：]?\s*", "摘要：", cleaned)
        normalized_labels += count
        cleaned, count = re.subn(r"^\s*(关键词|关键字)\s*[:：]?\s*", "关键词：", cleaned)
        normalized_labels += count
        cleaned, count = re.subn(r"^\s*(abstract)\s*[:：]?\s*", "Abstract: ", cleaned, flags=re.IGNORECASE)
        normalized_labels += count

        if cleaned.startswith("关键词：") and is_placeholder_keyword_value(cleaned.removeprefix("关键词：")):
            cleaned = "关键词：" + "；".join(keyword_suggestion)

        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned).strip()
        cleaned = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", cleaned)

        if cleaned == original:
            continue
        if cleaned:
            replace_paragraph_text(paragraph, cleaned)
        else:
            delete_paragraph(paragraph)

    if removed_placeholders:
        changes.append(f"已清理 {removed_placeholders} 处模板占位文字。")
    if removed_style_notes:
        changes.append(f"已删除 {removed_style_notes} 处模板格式说明。")
    if removed_template_codes:
        changes.append(f"已清理 {removed_template_codes} 处异常模板编号。")
    if normalized_labels:
        changes.append(f"已规范 {normalized_labels} 处摘要、关键词或英文摘要标签。")
    return changes


def split_mixed_heading_paragraphs(document) -> list[str]:
    split_count = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        inline_parts = split_inline_numbered_heading_text(text)
        if inline_parts:
            replace_paragraph_text(paragraph, inline_parts[0])
            anchor = paragraph
            for part in inline_parts[1:]:
                anchor = insert_paragraph_after(anchor, part)
            split_count += 1
            continue

        result = split_mixed_heading_text(text)
        if not result:
            continue
        heading, body = result
        replace_paragraph_text(paragraph, heading)
        insert_paragraph_after(paragraph, body)
        split_count += 1
    return [f"已拆分 {split_count} 个标题正文混排段落。"] if split_count else []


def split_inline_numbered_heading_text(text: str) -> list[str] | None:
    compact = text.strip()
    if len(compact) < 20:
        return None

    match = re.search(
        r"(?<!^)(?P<heading>\d{1,2}[.．、]\s*(?:引言|绪论|结论|结语|总结|参考文献|研究背景|研究方法|实验结果|结果与分析|讨论|致谢))\s*[：:]\s*(?P<body>\S.{5,})$",
        compact,
    )
    if not match:
        return None

    prefix = compact[: match.start()].strip()
    heading = normalize_numbered_heading_spacing(match.group("heading").strip())
    body = match.group("body").strip()
    if not prefix or not body:
        return None
    return [prefix, heading, body]


def normalize_numbered_heading_spacing(heading: str) -> str:
    return re.sub(r"^(\d{1,2})([.．、])\s*", r"\1\2 ", heading).strip()


def apply_page_setup(document, profile: dict[str, Any] | None) -> None:
    margins = (profile or {}).get("margins", {})
    for section in document.sections:
        section.top_margin = Cm(safe_margin(margins.get("top"), 2.54))
        section.bottom_margin = Cm(safe_margin(margins.get("bottom"), 2.54))
        section.left_margin = Cm(safe_margin(margins.get("left"), 3.18))
        section.right_margin = Cm(safe_margin(margins.get("right"), 3.18))


def apply_normal_style(document, profile: dict[str, Any] | None) -> None:
    normal_profile = (profile or {}).get("normal", {})
    font_name = normal_profile.get("font_name") or "宋体"
    font_size = safe_size(normal_profile.get("font_size"), 12) or 12
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_paragraph_styles(document, profile: dict[str, Any] | None) -> None:
    body_profile = (profile or {}).get("body", {})
    title_done = False
    in_references = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        clear_run_format(paragraph, profile)

        if not title_done:
            paragraph.style = get_or_create_style(document, "论文标题")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_style(paragraph, size=16, bold=True, font_name="黑体")
            title_done = True
            continue

        if text in {"参考文献", "References"} or text.lower() == "references":
            in_references = True
            format_heading(document, paragraph, profile, 1)
            continue

        if is_section_heading(text):
            format_heading(document, paragraph, profile, guess_heading_level(text))
            continue

        paragraph.alignment = body_profile.get("alignment") or WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = safe_spacing(body_profile.get("line_spacing"), 1.5)
        paragraph.paragraph_format.space_before = Pt(safe_size(body_profile.get("space_before_pt"), 0) or 0)
        paragraph.paragraph_format.space_after = Pt(safe_size(body_profile.get("space_after_pt"), 0) or 0)
        indent = safe_margin(body_profile.get("first_line_indent_cm"), 0.74)
        paragraph.paragraph_format.first_line_indent = None if in_references else Cm(indent)


def apply_table_styles(document) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.0


def format_heading(document, paragraph, profile: dict[str, Any] | None, level: int) -> None:
    paragraph.style = find_heading_style(document, level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    heading_profile = find_heading_profile(profile, level)
    set_run_style(
        paragraph,
        size=safe_size(heading_profile.get("font_size"), 14 if level == 1 else 12) or (14 if level == 1 else 12),
        bold=heading_profile.get("bold", True),
        font_name=heading_profile.get("font_name") or "黑体",
        color=RGBColor(0, 0, 0),
    )


def find_heading_style(document, level: int):
    for style_name in [f"Heading {level}", f"标题 {level}"]:
        try:
            return document.styles[style_name]
        except Exception:
            continue
    return get_or_create_style(document, f"标题 {level}")


def get_or_create_style(document, style_name: str):
    try:
        return document.styles[style_name]
    except Exception:
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        return style


def find_heading_profile(profile: dict[str, Any] | None, level: int) -> dict[str, Any]:
    if not profile:
        return {}
    headings = profile.get("heading", {})
    for name in [f"Heading {level}", f"标题 {level}"]:
        if name in headings:
            return headings[name]
    return {}


def clear_run_format(paragraph, profile: dict[str, Any] | None) -> None:
    normal = (profile or {}).get("normal", {})
    font_name = normal.get("font_name") or "宋体"
    font_size = safe_size(normal.get("font_size"), 12) or 12
    for run in paragraph.runs:
        set_run_font(run, font_name)
        run.font.size = Pt(font_size)
        run.bold = False


def set_run_style(paragraph, size: int | float, bold: bool | None, font_name: str | None = None, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        if font_name:
            set_run_font(run, font_name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def split_long_paragraphs(document, max_chars: int = 420) -> list[str]:
    split_count = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if len(text) <= max_chars or is_section_heading(text):
            continue
        chunks = split_text_by_sentence(text, max_chars)
        if len(chunks) <= 1:
            continue
        replace_paragraph_text(paragraph, chunks[0])
        anchor = paragraph
        for chunk in chunks[1:]:
            anchor = insert_paragraph_after(anchor, chunk)
        split_count += 1
    return [f"已按句意拆分 {split_count} 个过长段落。"] if split_count else []


def split_text_by_sentence(text: str, max_chars: int) -> list[str]:
    sentences = [item for item in re.split(r"(?<=[。！？!?])\s*", text) if item]
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        if current and len(current) + len(sentence) > max_chars:
            chunks.append(current.strip())
            current = sentence
        else:
            current += sentence
    if current.strip():
        chunks.append(current.strip())
    return chunks or [text]


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_element = deepcopy(paragraph._p)
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    replace_paragraph_text(new_paragraph, text)
    new_paragraph.style = paragraph.style
    return new_paragraph


def is_placeholder_keyword_value(value: str) -> bool:
    compact = re.sub(r"[\s；;,，、]+", "", value).lower()
    return not compact or any(marker in compact for marker in ["关键词", "关键字", "keyword", "xxx"]) or bool(re.fullmatch(r"[xX\d]+", compact))


def suggest_keywords(text: str) -> list[str]:
    rules = [
        ("人工智能", "人工智能"),
        ("论文", "论文写作"),
        ("教育", "教育研究"),
        ("心理", "心理健康"),
        ("压力", "压力管理"),
        ("学习", "学习效率"),
        ("格式", "格式规范"),
    ]
    keywords: list[str] = []
    for marker, keyword in rules:
        if marker in text and keyword not in keywords:
            keywords.append(keyword)
    return (keywords or ["研究主题", "问题分析", "优化策略"])[:5]
