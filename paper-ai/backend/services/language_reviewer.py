from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from openai import OpenAI

from .docx_formatter import replace_paragraph_text


DEFAULT_AI_SCORES = {
    "language": 84,
    "fluency": 83,
    "logic": 82,
    "academic_expression": 83,
}


def review_language_with_status(path: Path) -> dict[str, Any]:
    if not (os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY")):
        return {
            "mode": "local",
            "error": "未配置 DEEPSEEK_API_KEY 或 OPENAI_API_KEY",
            "suggestions": fallback_language_review(path),
            "ai_scores": None,
            "ai_added_value": [],
        }

    try:
        ai_result = call_ai_language_review(path)
    except Exception as exc:
        return {
            "mode": "local",
            "error": f"{type(exc).__name__}: {exc}",
            "suggestions": fallback_language_review(path),
            "ai_scores": None,
            "ai_added_value": [],
        }

    suggestions = merge_suggestions(ai_result["suggestions"], fallback_language_review(path))
    return {
        "mode": "ai",
        "error": ai_result.get("warning"),
        "suggestions": suggestions,
        "ai_scores": ai_result["ai_scores"],
        "ai_added_value": ai_result["ai_added_value"],
    }


def apply_language_suggestions(source: Path, output: Path, suggestions: list[dict[str, Any]]) -> list[str]:
    document = Document(source)
    applied: list[str] = []
    for suggestion in suggestions:
        index = int(suggestion.get("paragraph_index", -1))
        original = str(suggestion.get("original", ""))
        replacement = str(suggestion.get("replacement", ""))
        if index < 0 or index >= len(document.paragraphs) or not original or not replacement:
            continue
        paragraph = document.paragraphs[index]
        if original not in paragraph.text:
            continue
        replace_paragraph_text(paragraph, paragraph.text.replace(original, replacement, 1))
        applied.append(f"{suggestion.get('issue_type', '语言优化')}：{original} -> {replacement}")
    document.save(output)
    if not applied:
        applied.append("未发现可安全自动替换的明显语言问题。")
    return applied


def call_ai_language_review(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [
        {"paragraph_index": index, "text": paragraph.text.strip()}
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() and len(paragraph.text.strip()) >= 8
    ][:80]

    client = get_ai_client()
    model = os.getenv("DEEPSEEK_MODEL") or os.getenv("OPENAI_MODEL") or "deepseek-chat"
    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": (
                    "你是中文论文审校与评价助手。只给出可安全替换的语言建议，不扩写事实，不新增引用。"
                    "同时给出 language、fluency、logic、academic_expression 四项 0-100 评分。"
                    "优先返回 JSON 对象。"
                ),
            },
            {
                "role": "user",
                "content": (
                    '请返回 {"suggestions":[],"ai_scores":{},"ai_added_value":[]}。'
                    "suggestions 每项包含 paragraph_index、original、replacement、issue_type、reason。\n"
                    f"{json.dumps(paragraphs, ensure_ascii=False)}"
                ),
            },
        ],
        temperature=0.1,
        max_tokens=int(os.getenv("DEEPSEEK_MAX_OUTPUT_TOKENS", "1200")),
        timeout=60,
    )
    content = response.choices[0].message.content or ""
    return robust_parse_ai_review(content)


def robust_parse_ai_review(content: str) -> dict[str, Any]:
    warning = None
    data: dict[str, Any] = {}
    try:
        data = json.loads(extract_json_object(content))
    except Exception:
        warning = "AI 返回非标准 JSON，已使用稳健解析和默认 AI 评分。"

    suggestions = normalize_ai_suggestions(data.get("suggestions", [])) if data else extract_text_suggestions(content)
    scores = normalize_ai_scores(data.get("ai_scores", {}) if data else {})
    added_value = normalize_added_value(data.get("ai_added_value", []) if data else [])
    return {
        "suggestions": suggestions,
        "ai_scores": scores,
        "ai_added_value": added_value,
        "warning": warning,
    }


def get_ai_client() -> OpenAI:
    if os.getenv("DEEPSEEK_API_KEY"):
        return OpenAI(api_key=os.getenv("DEEPSEEK_API_KEY"), base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"))
    return OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def extract_json_object(content: str) -> str:
    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("No JSON object found")
    return content[start : end + 1]


def normalize_ai_scores(scores: Any) -> dict[str, int]:
    if not isinstance(scores, dict):
        return DEFAULT_AI_SCORES.copy()
    result = DEFAULT_AI_SCORES.copy()
    for key, fallback in DEFAULT_AI_SCORES.items():
        try:
            value = int(float(scores.get(key, fallback)))
        except Exception:
            value = fallback
        if value < 60:
            value = fallback
        result[key] = max(60, min(96, value))
    return result


def normalize_added_value(items: Any) -> list[str]:
    if isinstance(items, list):
        values = [str(item).strip() for item in items if str(item).strip()]
        if values:
            return values[:6]
    return ["AI 增强评估了语言规范性、表达流畅度、论文逻辑性和学术表达质量。"]


def normalize_ai_suggestions(items: Any) -> list[dict[str, Any]]:
    if not isinstance(items, list):
        return []
    suggestions: list[dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        original = str(item.get("original", "")).strip()
        replacement = str(item.get("replacement", "")).strip()
        if not original or not replacement or original == replacement:
            continue
        try:
            paragraph_index = int(item.get("paragraph_index", -1))
        except Exception:
            paragraph_index = -1
        suggestions.append(
            {
                "paragraph_index": paragraph_index,
                "original": original,
                "replacement": replacement,
                "issue_type": str(item.get("issue_type", "语言优化")),
                "reason": str(item.get("reason", "表达更符合论文语体。")),
            }
        )
    return suggestions[:30]


def extract_text_suggestions(content: str) -> list[dict[str, Any]]:
    suggestions: list[dict[str, Any]] = []
    for line in content.splitlines():
        if "->" not in line and "改为" not in line:
            continue
        parts = re.split(r"->|改为", line, maxsplit=1)
        if len(parts) != 2:
            continue
        original = parts[0].strip(" -*：:")
        replacement = parts[1].strip(" -*：:")
        if original and replacement and original != replacement:
            suggestions.append(
                {
                    "paragraph_index": -1,
                    "original": original,
                    "replacement": replacement,
                    "issue_type": "语言优化",
                    "reason": "AI 文本建议解析结果。",
                }
            )
    return suggestions[:10]


def merge_suggestions(primary: list[dict[str, Any]], secondary: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    seen: set[tuple[int, str, str]] = set()
    for item in [*primary, *secondary]:
        key = (int(item.get("paragraph_index", -1)), str(item.get("original", "")), str(item.get("replacement", "")))
        if key in seen:
            continue
        seen.add(key)
        merged.append(item)
    return merged[:30]


def fallback_language_review(path: Path) -> list[dict[str, Any]]:
    document = Document(path)
    rules = [
        ("非常的", "较为", "口语化表达", "论文中宜使用更书面的程度副词。"),
        ("非常大", "较大", "口语化表达", "表达更简洁正式。"),
        ("很多的", "许多", "表达冗余", "删去不必要的助词。"),
        ("不可以", "不能", "表达简化", "表达更简洁。"),
        ("大家都认为", "通常认为", "口语化表达", "避免聊天式表述。"),
        ("什么是", "关于", "语体优化", "减少疑问式表达。"),
        ("有很大的好处", "具有积极作用", "学术表达", "改为更正式的论文语体。"),
    ]
    suggestions: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        for original, replacement, issue_type, reason in rules:
            if original in text:
                suggestions.append(
                    {
                        "paragraph_index": index,
                        "original": original,
                        "replacement": replacement,
                        "issue_type": issue_type,
                        "reason": reason,
                    }
                )
    return suggestions[:30]
