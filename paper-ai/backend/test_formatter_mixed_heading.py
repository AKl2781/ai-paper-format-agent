from pathlib import Path
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from services.docx_formatter import apply_paper_format


def test_split_inline_numbered_heading(tmp_path: Path) -> None:
    source = tmp_path / "inline_heading.docx"
    output = tmp_path / "inline_heading_formatted.docx"

    document = Document()
    document.add_paragraph("测试论文")
    document.add_paragraph(
        "前文介绍了一些压力的表现以及办法。4. 结语：面对高中生的学习压力，文中介绍了压力形成的原因。"
    )
    document.save(source)

    apply_paper_format(source, output)

    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs if paragraph.text.strip()]
    assert "前文介绍了一些压力的表现以及办法。" in paragraphs
    assert "4. 结语" in paragraphs
    assert "面对高中生的学习压力，文中介绍了压力形成的原因。" in paragraphs


def test_cover_page_preserved_and_template_cover_not_used_as_body(tmp_path: Path) -> None:
    source = tmp_path / "course_paper.docx"
    template = tmp_path / "course_template.docx"
    output = tmp_path / "course_paper_formatted.docx"

    document = Document()
    cover_title = document.add_paragraph("课 程 论 文")
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text in [
        "课\u3000程：《果蔬营养与保健》",
        "姓\u3000名：__________",
        "学\u3000号：__________",
        "学\u3000院：生物与农业学院",
        "专\u3000业：园艺",
        "班\u3000级：__________",
    ]:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.left_indent = Cm(3)
    document.add_paragraph("（教务部制表）").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    document.add_paragraph("荔枝品种特性、营养成分与食药保健价值及产业化应用研究")
    document.add_paragraph("摘要：本文围绕荔枝营养价值展开分析。")
    document.add_paragraph("荔枝含有糖类、维生素和多酚等成分，具有较高的食用价值。")
    document.save(source)

    template_doc = Document()
    template_doc.add_paragraph("课 程 论 文").alignment = WD_ALIGN_PARAGRAPH.CENTER
    template_doc.add_paragraph("姓\u3000名：__________").alignment = WD_ALIGN_PARAGRAPH.CENTER
    template_doc.add_page_break()
    template_doc.save(template)

    apply_paper_format(source, output, template)

    formatted = Document(output)
    paragraphs = formatted.paragraphs
    assert paragraphs[0].text == "课 程 论 文"
    assert paragraphs[1].text == "课\u3000程：《果蔬营养与保健》"
    assert abs(paragraphs[1].paragraph_format.left_indent.cm - 3) < 0.01

    body = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("荔枝含有"))
    assert body.alignment != WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    with tempfile.TemporaryDirectory() as directory:
        tmp_path = Path(directory)
        test_split_inline_numbered_heading(tmp_path)
        test_cover_page_preserved_and_template_cover_not_used_as_body(tmp_path)
    print("test_formatter_mixed_heading PASS")
