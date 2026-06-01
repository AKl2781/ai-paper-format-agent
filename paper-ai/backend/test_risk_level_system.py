from __future__ import annotations

from docx import Document

from services.agent_orchestrator import needs_manual_review
from services.docx_analyzer import analyze_figure_tables, analyze_references


def make_document(paragraphs: list[str], *, table_count: int = 0) -> Document:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    for _ in range(table_count):
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "item"
        table.cell(0, 1).text = "value"
    return document


def assert_equal(name: str, actual: object, expected: object) -> None:
    if actual != expected:
        raise AssertionError(f"{name} FAIL expected={expected!r} actual={actual!r}")
    print(f"{name} PASS")


def levels(check: dict[str, object]) -> list[str]:
    return [item["level"] for item in check.get("risk_items", [])]


def has_level(check: dict[str, object], level: str) -> bool:
    return level in levels(check)


def check_figure_tables(paragraphs: list[str], *, table_count: int = 0) -> dict[str, object]:
    document = make_document(paragraphs, table_count=table_count)
    return analyze_figure_tables([p.text.strip() for p in document.paragraphs if p.text.strip()], document)


def main() -> None:
    no_reference = analyze_references(["Paper title", "Body without references."])
    assert_equal("missing_reference_section_level", levels(no_reference), ["warning"])
    assert_equal("missing_reference_section_manual_review", needs_manual_review(None, {"reference_check": no_reference}, False), False)

    duplicate_reference = analyze_references(["Body cites [1].", "References", "[1] A.", "[1] B."])
    assert_equal("duplicate_reference_high_risk", has_level(duplicate_reference, "high_risk"), True)
    assert_equal("duplicate_reference_manual_review", needs_manual_review(None, {"reference_check": duplicate_reference}, False), True)

    missing_reference = analyze_references(["Body cites [2].", "References", "[1] A."])
    assert_equal("missing_reference_high_risk", has_level(missing_reference, "high_risk"), True)

    uncited_reference = analyze_references(["Body cites [1].", "References", "[1] A.", "[2] B."])
    assert_equal("uncited_reference_warning", levels(uncited_reference), ["warning"])

    figure_gap = check_figure_tables(["Figure 1 flow", "Figure 3 result"])
    assert_equal("figure_gap_warning", levels(figure_gap), ["warning"])
    assert_equal("figure_gap_manual_review", needs_manual_review(None, {"figure_table_check": figure_gap}, False), False)

    duplicate_table = check_figure_tables(["Table 1 sample", "Table 1 result"])
    assert_equal("duplicate_table_high_risk", has_level(duplicate_table, "high_risk"), True)

    missing_figure_reference = check_figure_tables(["As Figure 2 shows.", "Figure 1 flow"])
    assert_equal("missing_figure_reference_high_risk", has_level(missing_figure_reference, "high_risk"), True)

    table_without_caption = check_figure_tables(["This is a form-like report."], table_count=1)
    assert_equal("table_without_caption_info", levels(table_without_caption), ["info"])
    assert_equal("table_without_caption_manual_review", needs_manual_review(None, {"figure_table_check": table_without_caption}, False), False)

    assert_equal("requires_confirmation_blocking", needs_manual_review(None, None, True), True)
    print("RISK_LEVEL_SYSTEM PASS")


if __name__ == "__main__":
    main()
